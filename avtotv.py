import re
import typing
from striprtf.striprtf import rtf_to_text
from docx import Document

def _is_show(entry: str) -> bool:
    pattern = r'([0-9]{1,2}\.[0-9]{1,2}) (.+)'
    return re.search(pattern, entry)


def _is_channel(entry: str):
    pattern = r'\w+'
    return re.search(pattern, entry)


def _is_date(entry: str):
    pattern = r'\w+\, [0-9]{1,2} \w+'
    return re.search(pattern, entry)


def main():
    doc = Document()
    list_paragraph = doc.add_paragraph()
    with open('1.rtf') as f:
        entries = re.split("\n+", f.read())
        for entry in entries:
            entry = rtf_to_text(entry).strip('/n')
            if _is_show(entry):
                list_paragraph.add_run(_is_show(entry).group(1)).bold = True
                list_paragraph.add_run('\t')
                list_paragraph.add_run(_is_show(entry).group(2))
                list_paragraph.add_run('\n')
                print(_is_show(entry))
            elif _is_date(entry):
                
                print(_is_date(entry).group(0))
            elif _is_channel(entry):
                print(_is_channel(entry).group(0))
    doc.save('res.docx')

if __name__ == '__main__':
    main()